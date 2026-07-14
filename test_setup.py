from docx import Document
from fpdf import FPDF

# Test python-docx
doc = Document()
doc.add_heading("SentinelAI Test", level=1)
doc.add_paragraph("If you can see this, python-docx works.")
doc.save("test_report.docx")

# Test fpdf2
pdf = FPDF()
pdf.add_page()
pdf.set_font("Helvetica", size=12)
pdf.cell(200, 10, text="If you can see this, fpdf2 works.", new_x="LMARGIN", new_y="NEXT")
pdf.output("test_report.pdf")

print("Environment confirmed working: docx and pdf generated successfully.")